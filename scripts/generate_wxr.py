#!/usr/bin/env python3
"""Create WordPress WXR draft imports from Chinese Word articles without rewriting text."""

from __future__ import annotations

import argparse
import csv
import html
import io
import json
import os
import re
import secrets
import sys
import tempfile
import zipfile
from collections import Counter
from dataclasses import dataclass
from datetime import datetime, timezone
from email.utils import format_datetime
from pathlib import Path
from typing import Iterable
from urllib.parse import quote, urlsplit, urlunsplit

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


SOLUTION_TOPICS = (
    "style3d",
    "style 3d",
    "凌迪",
    "ai服装",
    "ai 时尚",
    "3d+ai服装",
    "3d＋ai服装",
    "3d ai服装",
    "3d+ai 时尚",
)
SOLUTION_MARKERS = ("解决方案", "方案", "落地", "实施", "部署", "建设", "升级", "改造", "怎么建")
KNOWLEDGE_MARKERS = ("是什么", "有哪些", "定义", "特性", "原理", "概述", "详解", "教程", "怎么用")
NUMERIC_SUBHEADING = re.compile(r"^\s*(?:\d+|[一二三四五六七八九十百]+)\s*(?:[.．、)）]\s*).+\S\s*$")
FILENAME = re.compile(r"^([A-Za-z]+-\d+-\d+)[.．、]\s*(.+)$")
META_FIELDS = {
    "focuskw": re.compile(r"^\s*(?:焦点关键词|关键词|主关键词|focus\s*keyword)\s*[:：]\s*(.+?)\s*$", re.I),
    "seo_title": re.compile(r"^\s*(?:seo\s*标题|seo\s*title|meta\s*title|标题标签)\s*[:：]\s*(.+?)\s*$", re.I),
    "metadesc": re.compile(r"^\s*(?:seo\s*描述|meta\s*description|描述|摘要)\s*[:：]\s*(.+?)\s*$", re.I),
}


BODY_MARKERS = ("\u6b63\u6587\uff1a", "\u6b63\u6587:")
DEFAULT_COMPETITORS = (
    "CLO 3D", "CLO Virtual Fashion", "Marvelous Designer", "Browzwear",
    "VStitcher", "Lotta", "Stylezone", "Optitex", "Lectra", "Modaris",
    "Gerber", "AccuMark", "TUKAcad", "TUKA3D", "Audaces", "Assyst", "Vidya",
)
STYLE3D_FRAME_IMAGE_URLS = tuple(
    f"https://help-zh.style3d.com/wp-content/uploads/2026/08/Frame-{number}.png"
    for number in range(1, 54)
)


@dataclass
class SourceDocument:
    name: str
    data: bytes


@dataclass
class Article:
    article_id: str
    title: str
    content_html: str
    category: str
    category_reason: str
    focus_keyword: str
    seo_title: str
    meta_description: str
    source: str
    image_count: int
    seo_warnings: tuple[str, ...]


@dataclass
class ContentAudit:
    article_id: str
    title: str
    source: str
    matches: tuple[tuple[str, str], ...]


@dataclass(frozen=True)
class IntroImageAssignment:
    article_id: str
    title: str
    url: str
    cycle: int


def cdata(value: str) -> str:
    return "<![CDATA[" + value.replace("]]>", "]]]]><![CDATA[>") + "]]>"


def clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def source_title(name: str) -> tuple[str, str]:
    stem = Path(name).stem
    match = FILENAME.match(stem)
    if match:
        return match.group(1), clean_text(match.group(2))
    safe_id = re.sub(r"[^A-Za-z0-9-]+", "-", stem).strip("-").lower() or "article"
    return safe_id, clean_text(stem)


def iter_sources(source: Path) -> Iterable[SourceDocument]:
    if source.is_file() and source.suffix.lower() == ".docx":
        yield SourceDocument(source.name, source.read_bytes())
        return
    if source.is_file() and source.suffix.lower() == ".zip":
        archives = [source]
    elif source.is_dir():
        archives = sorted(source.rglob("*.zip"))
        for docx_path in sorted(source.rglob("*.docx")):
            yield SourceDocument(str(docx_path.relative_to(source)), docx_path.read_bytes())
    else:
        raise ValueError(f"输入不存在或格式不支持: {source}")

    for archive_path in archives:
        with zipfile.ZipFile(archive_path) as archive:
            for entry in sorted(archive.infolist(), key=lambda item: item.filename.casefold()):
                if entry.is_dir() or not entry.filename.lower().endswith(".docx"):
                    continue
                yield SourceDocument(f"{archive_path.name}!{entry.filename}", archive.read(entry))


def load_competitors(path: str | None) -> tuple[str, ...]:
    if not path:
        return DEFAULT_COMPETITORS
    payload = json.loads(Path(path).read_text(encoding="utf-8"))
    values = payload.get("competitors") if isinstance(payload, dict) else payload
    if not isinstance(values, list) or not all(isinstance(value, str) and clean_text(value) for value in values):
        raise ValueError("Competitor list must be a JSON array of nonempty names or an object with a competitors array")
    return tuple(dict.fromkeys(clean_text(value) for value in values))


def load_intro_image_pool(path: str | None, use_style3d_frames: bool) -> tuple[str, ...]:
    if path and use_style3d_frames:
        raise ValueError("Use either --intro-image-pool or --use-style3d-frame-pool, not both")
    if use_style3d_frames:
        return STYLE3D_FRAME_IMAGE_URLS
    if not path:
        return ()
    payload = json.loads(Path(path).read_text(encoding="utf-8"))
    values = payload.get("images") if isinstance(payload, dict) else payload
    if not isinstance(values, list) or not values:
        raise ValueError("Intro image pool must be a nonempty JSON array or an object with an images array")
    urls = tuple(clean_text(value) for value in values if isinstance(value, str))
    if len(urls) != len(values) or len(set(urls)) != len(urls):
        raise ValueError("Intro image URLs must be nonempty, unique strings")
    for url in urls:
        parsed = urlsplit(url)
        if parsed.scheme not in {"http", "https"} or not parsed.netloc:
            raise ValueError(f"Intro image URL is not an absolute HTTP(S) URL: {url}")
    return urls


def load_intro_image_state(path: Path, pool: tuple[str, ...]) -> tuple[list[str], int]:
    if not path.exists():
        return list(pool), 0
    payload = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict) or payload.get("version") != 1:
        raise ValueError("Intro image state must be a version 1 JSON object")
    remaining = payload.get("remaining")
    cycle = payload.get("cycle")
    if not isinstance(remaining, list) or not isinstance(cycle, int) or cycle < 0:
        raise ValueError("Intro image state must contain a nonnegative cycle and remaining image URLs")
    if len(remaining) != len(set(remaining)) or any(url not in pool for url in remaining):
        raise ValueError("Intro image state does not match the configured image pool")
    return list(remaining), cycle


def assign_intro_images(articles: list[Article], pool: tuple[str, ...], state_path: Path) -> tuple[list[IntroImageAssignment], dict[str, object]]:
    remaining, cycle = load_intro_image_state(state_path, pool)
    chooser = secrets.SystemRandom()
    assignments: list[IntroImageAssignment] = []
    for article in articles:
        if not remaining:
            remaining = list(pool)
            cycle += 1
        choice_index = chooser.randrange(len(remaining))
        url = remaining.pop(choice_index)
        assignments.append(IntroImageAssignment(article.article_id, article.title, url, cycle))
    return assignments, {"version": 1, "remaining": remaining, "cycle": cycle}


def render_intro_image(url: str, title: str) -> str:
    return f'<p><img src="{html.escape(url, quote=True)}" alt="{html.escape(title, quote=True)}" loading="lazy" /></p>'


def write_intro_image_state(path: Path, state: dict[str, object]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile("w", encoding="utf-8", dir=path.parent, delete=False, suffix=".tmp") as handle:
        json.dump(state, handle, ensure_ascii=False, indent=2)
        handle.write("\n")
        temp_path = Path(handle.name)
    try:
        os.replace(temp_path, path)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def write_intro_image_report(assignments: list[IntroImageAssignment], path: Path) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=("article_id", "title", "intro_image_url", "cycle"))
        writer.writeheader()
        for assignment in assignments:
            writer.writerow({
                "article_id": assignment.article_id,
                "title": assignment.title,
                "intro_image_url": assignment.url,
                "cycle": assignment.cycle,
            })


def document_text(document: Document) -> str:
    blocks = [paragraph.text for paragraph in document.paragraphs if clean_text(paragraph.text)]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                blocks.extend(paragraph.text for paragraph in cell.paragraphs if clean_text(paragraph.text))
    return "\n".join(blocks)


def audit_source(source: SourceDocument, competitors: tuple[str, ...]) -> ContentAudit:
    article_id, title = source_title(source.name.rsplit("!", 1)[-1])
    raw_text = document_text(Document(io.BytesIO(source.data)))
    normalized = raw_text.casefold()
    matches: list[tuple[str, str]] = []
    for competitor in competitors:
        start = normalized.find(competitor.casefold())
        if start == -1:
            continue
        context = clean_text(raw_text[max(0, start - 60) : start + len(competitor) + 100])
        matches.append((competitor, context))
    return ContentAudit(article_id, title, source.name, tuple(matches))


def write_content_audit(audits: list[ContentAudit], path: Path) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=("article_id", "title", "status", "competitor", "context", "source"))
        writer.writeheader()
        for audit in audits:
            if not audit.matches:
                writer.writerow({"article_id": audit.article_id, "title": audit.title, "status": "No configured competitor name found; semantic review still required", "competitor": "", "context": "", "source": audit.source})
            for competitor, context in audit.matches:
                writer.writerow({"article_id": audit.article_id, "title": audit.title, "status": "Blocked: competitor reference found", "competitor": competitor, "context": context, "source": audit.source})


def metadata_from_document(document: Document) -> dict[str, str]:
    metadata: dict[str, str] = {}
    # Only accept metadata before the optional body marker.
    for child in document.element.body.iterchildren():
        if child.tag.rsplit("}", 1)[-1] != "p":
            continue
        paragraph = Paragraph(child, document)
        text = clean_text(paragraph.text)
        if is_body_marker(text):
            break
        for key, pattern in META_FIELDS.items():
            match = pattern.match(text)
            if match and key not in metadata:
                metadata[key] = match.group(1)
    return metadata

def is_metadata_line(text: str) -> bool:
    clean = clean_text(text)
    return any(pattern.match(clean) for pattern in META_FIELDS.values())


def is_body_marker(text: str) -> bool:
    return clean_text(text) in BODY_MARKERS


def body_children(document: Document):
    children = list(document.element.body.iterchildren())
    for index, child in enumerate(children):
        if child.tag.rsplit("}", 1)[-1] != "p":
            continue
        if is_body_marker(Paragraph(child, document).text):
            return children[index + 1 :]
    return children


def escape_run_text(value: str) -> str:
    return html.escape(value).replace("\t", "&emsp;").replace("\n", "<br>")


def paragraph_inline_html(paragraph: Paragraph, force_plain: bool = False) -> str:
    if force_plain or "".join(run.text for run in paragraph.runs) != paragraph.text:
        return escape_run_text(paragraph.text)
    parts: list[str] = []
    for run in paragraph.runs:
        text = escape_run_text(run.text)
        if not text:
            continue
        if run.bold:
            text = f"<strong>{text}</strong>"
        if run.italic:
            text = f"<em>{text}</em>"
        if run.underline:
            text = f"<u>{text}</u>"
        parts.append(text)
    return "".join(parts) or escape_run_text(paragraph.text)


def heading_tag(paragraph: Paragraph) -> str | None:
    style = (paragraph.style.name if paragraph.style else "").casefold()
    match = re.search(r"(?:heading|标题)\s*([1-6])", style)
    if not match:
        return None
    return "h2"


def list_kind(paragraph: Paragraph) -> str | None:
    style = (paragraph.style.name if paragraph.style else "").casefold()
    if "bullet" in style or "项目符号" in style:
        return "ul"
    if "list" in style or "number" in style or "编号" in style:
        return "ol"
    return None


def table_html(table: Table) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            value = "<br>".join(paragraph_inline_html(paragraph) for paragraph in cell.paragraphs if paragraph.text)
            cells.append(f"<td>{value}</td>")
        rows.append("<tr>" + "".join(cells) + "</tr>")
    return "<table><tbody>" + "".join(rows) + "</tbody></table>"


def render_body(document: Document, title: str) -> str:
    blocks: list[str] = []
    pending_list: str | None = None
    pending_items: list[str] = []

    def flush_list() -> None:
        nonlocal pending_list, pending_items
        if pending_list:
            blocks.append(f"<{pending_list}>" + "".join(pending_items) + f"</{pending_list}>")
        pending_list = None
        pending_items = []

    for child in body_children(document):
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "tbl":
            flush_list()
            blocks.append(table_html(Table(child, document)))
            continue
        if tag != "p":
            continue
        paragraph = Paragraph(child, document)
        raw = paragraph.text
        text = clean_text(raw)
        if not text or is_metadata_line(raw) or text == clean_text(title):
            continue
        if NUMERIC_SUBHEADING.match(text):
            flush_list()
            blocks.append(f'<p style="font-size: 16px;"><strong>{paragraph_inline_html(paragraph, force_plain=True)}</strong></p>')
            continue
        if heading_tag(paragraph):
            flush_list()
            blocks.append(f'<h2 style="font-size: 20px; font-weight: normal;">{paragraph_inline_html(paragraph, force_plain=True)}</h2>')
            continue
        kind = list_kind(paragraph)
        if kind:
            if pending_list and pending_list != kind:
                flush_list()
            pending_list = kind
            pending_items.append(f"<li>{paragraph_inline_html(paragraph)}</li>")
            continue
        flush_list()
        blocks.append(f"<p>{paragraph_inline_html(paragraph)}</p>")
    flush_list()
    return "\n".join(blocks)

def derive_focus_keyword(title: str) -> str:
    title = clean_text(title)
    match = re.match(r"^什么是(.+?)[？?]", title)
    if match:
        return match.group(1).strip()
    match = re.match(r"^(.+?)(?:怎么|如何|是什么|有哪些|为什么|能否)", title)
    if match:
        return match.group(1).strip()
    return re.split(r"[？?：:，,。]", title, maxsplit=1)[0].strip()


def derive_seo_title(title: str, keyword: str, brand_suffix: str) -> str:
    question = re.split(r"[？?]", clean_text(title), maxsplit=1)[0]
    if "？" in title:
        question += "？"
    elif "?" in title:
        question += "?"
    return f"{question}_{keyword}-{brand_suffix}"


def classify(title: str, article_id: str, overrides: dict[str, str]) -> tuple[str, str]:
    if article_id in overrides:
        return overrides[article_id], "人工确认"
    normalized = title.casefold().replace("＋", "+")
    if any(topic in normalized for topic in SOLUTION_TOPICS):
        return "解决方案", "AI/3D+AI/Style3D 主题"
    has_solution = any(marker.casefold() in normalized for marker in SOLUTION_MARKERS)
    has_knowledge = any(marker.casefold() in normalized for marker in KNOWLEDGE_MARKERS)
    if has_solution and not has_knowledge:
        return "解决方案", "明确方案/实施主题"
    if has_knowledge and not has_solution:
        return "行业知识", "明确知识科普主题"
    return "待确认", "标题无法按既定规则可靠判定"


def first_body_text(document: Document, title: str) -> str:
    for child in body_children(document):
        if child.tag.rsplit("}", 1)[-1] != "p":
            continue
        paragraph = Paragraph(child, document)
        text = clean_text(paragraph.text)
        if not text or is_body_marker(text) or text == clean_text(title) or is_metadata_line(text):
            continue
        if NUMERIC_SUBHEADING.match(text) or heading_tag(paragraph):
            continue
        return text
    return ""


def seo_warnings(keyword: str, seo_title: str, meta_description: str) -> tuple[str, ...]:
    warnings: list[str] = []
    if not keyword:
        warnings.append("focus keyword is empty")
    if not seo_title:
        warnings.append("SEO title is empty")
    if not meta_description:
        warnings.append("SEO description is empty")
    if seo_title and not 15 <= len(seo_title) <= 60:
        warnings.append(f"SEO title length is {len(seo_title)} (recommended: 15-60)")
    if meta_description and not 50 <= len(meta_description) <= 160:
        warnings.append(f"SEO description length is {len(meta_description)} (recommended: 50-160)")
    if keyword and seo_title and keyword.casefold() not in seo_title.casefold():
        warnings.append("focus keyword is missing from SEO title")
    if keyword and meta_description and keyword.casefold() not in meta_description.casefold():
        warnings.append("focus keyword is missing from SEO description")
    return tuple(warnings)

def make_article(source: SourceDocument, overrides: dict[str, str], brand_suffix: str) -> Article:
    document = Document(io.BytesIO(source.data))
    article_id, title = source_title(source.name.rsplit("!", 1)[-1])
    metadata = metadata_from_document(document)
    keyword = metadata.get("focuskw") or derive_focus_keyword(title)
    meta_description = metadata.get("metadesc") or first_body_text(document, title)
    seo_title = metadata.get("seo_title") or derive_seo_title(title, keyword, brand_suffix)
    category, reason = classify(title, article_id, overrides)
    return Article(
        article_id=article_id,
        title=title,
        content_html=render_body(document, title),
        category=category,
        category_reason=reason,
        focus_keyword=keyword,
        seo_title=seo_title,
        meta_description=meta_description,
        source=source.name,
        image_count=len(document.inline_shapes),
        seo_warnings=seo_warnings(keyword, seo_title, meta_description),
    )


def wp_date(now: datetime) -> tuple[str, str]:
    utc = now.astimezone(timezone.utc)
    return now.strftime("%Y-%m-%d %H:%M:%S"), utc.strftime("%Y-%m-%d %H:%M:%S")


def write_wxr(
    articles: list[Article],
    output: Path,
    args: argparse.Namespace,
    intro_images: dict[str, IntroImageAssignment],
) -> None:
    now = datetime.now().astimezone()
    local_date, utc_date = wp_date(now)
    pub_date = format_datetime(now.astimezone(timezone.utc), usegmt=True)
    parsed_url = urlsplit(args.site_url)
    base_site_url = urlunsplit((parsed_url.scheme, parsed_url.netloc, "", "", ""))
    base_blog_url = args.site_url.rstrip("/") + "/"
    categories = sorted({article.category for article in articles})
    tags = sorted({article.focus_keyword for article in articles if article.focus_keyword})
    lines = [
        '<?xml version="1.0" encoding="utf-8"?>',
        '<rss version="2.0" xmlns:excerpt="http://wordpress.org/export/1.2/excerpt/" xmlns:content="http://purl.org/rss/1.0/modules/content/" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:wp="http://wordpress.org/export/1.2/">',
        "  <channel>",
        f"    <title>{cdata('Style3D WordPress Draft Import')}</title>",
        f"    <link>{html.escape(base_blog_url)}</link>",
        f"    <description>{cdata(f'{len(articles)} draft articles prepared from Word files')}</description>",
        f"    <pubDate>{pub_date}</pubDate>",
        "    <language>zh-CN</language>",
        "    <wp:wxr_version>1.2</wp:wxr_version>",
        f"    <wp:base_site_url>{html.escape(base_site_url)}</wp:base_site_url>",
        f"    <wp:base_blog_url>{html.escape(base_blog_url)}</wp:base_blog_url>",
        "    <wp:author>",
        f"      <wp:author_id>{args.author_id}</wp:author_id>",
        f"      <wp:author_login>{cdata(args.author_login)}</wp:author_login>",
        "      <wp:author_email><![CDATA[]]></wp:author_email>",
        f"      <wp:author_display_name>{cdata(args.author_login)}</wp:author_display_name>",
        "      <wp:author_first_name><![CDATA[]]></wp:author_first_name>",
        "      <wp:author_last_name><![CDATA[]]></wp:author_last_name>",
        "    </wp:author>",
    ]
    for category in categories:
        lines.extend((
            "    <wp:category>",
            "      <wp:term_id>0</wp:term_id>",
            f"      <wp:category_nicename>{cdata(category)}</wp:category_nicename>",
            "      <wp:category_parent><![CDATA[]]></wp:category_parent>",
            f"      <wp:cat_name>{cdata(category)}</wp:cat_name>",
            "    </wp:category>",
        ))
    for tag in tags:
        lines.extend((
            "    <wp:tag>",
            "      <wp:term_id>0</wp:term_id>",
            f"      <wp:tag_slug>{cdata(tag)}</wp:tag_slug>",
            f"      <wp:tag_name>{cdata(tag)}</wp:tag_name>",
            "    </wp:tag>",
        ))
    for article in articles:
        link = f"{base_blog_url}?codex_import={quote(article.article_id)}"
        assignment = intro_images.get(article.article_id)
        content_html = article.content_html
        if assignment:
            content_html = render_intro_image(assignment.url, article.title) + "\n" + content_html
        lines.extend((
            "    <item>",
            f"      <title>{cdata(article.title)}</title>",
            f"      <link>{html.escape(link)}</link>",
            f"      <pubDate>{pub_date}</pubDate>",
            f"      <dc:creator>{cdata(args.author_login)}</dc:creator>",
            f"      <guid isPermaLink=\"false\">{html.escape('style3d-import-' + article.article_id)}</guid>",
            "      <description></description>",
            f"      <content:encoded>{cdata(content_html)}</content:encoded>",
            "      <excerpt:encoded><![CDATA[]]></excerpt:encoded>",
            "      <wp:post_id>0</wp:post_id>",
            f"      <wp:post_date>{cdata(local_date)}</wp:post_date>",
            f"      <wp:post_date_gmt>{cdata(utc_date)}</wp:post_date_gmt>",
            "      <wp:comment_status><![CDATA[closed]]></wp:comment_status>",
            "      <wp:ping_status><![CDATA[closed]]></wp:ping_status>",
            f"      <wp:post_name>{cdata(article.article_id.casefold())}</wp:post_name>",
            "      <wp:status><![CDATA[draft]]></wp:status>",
            "      <wp:post_parent>0</wp:post_parent>",
            "      <wp:menu_order>0</wp:menu_order>",
            "      <wp:post_type><![CDATA[post]]></wp:post_type>",
            "      <wp:post_password><![CDATA[]]></wp:post_password>",
            "      <wp:is_sticky>0</wp:is_sticky>",
            f"      <category domain=\"category\" nicename=\"{html.escape(article.category, quote=True)}\">{cdata(article.category)}</category>",
            f"      <category domain=\"post_tag\" nicename=\"{html.escape(article.focus_keyword, quote=True)}\">{cdata(article.focus_keyword)}</category>",
            "      <wp:postmeta>",
            "        <wp:meta_key><![CDATA[_yoast_wpseo_focuskw]]></wp:meta_key>",
            f"        <wp:meta_value>{cdata(article.focus_keyword)}</wp:meta_value>",
            "      </wp:postmeta>",
            "      <wp:postmeta>",
            "        <wp:meta_key><![CDATA[_yoast_wpseo_title]]></wp:meta_key>",
            f"        <wp:meta_value>{cdata(article.seo_title)}</wp:meta_value>",
            "      </wp:postmeta>",
            "      <wp:postmeta>",
            "        <wp:meta_key><![CDATA[_yoast_wpseo_metadesc]]></wp:meta_key>",
            f"        <wp:meta_value>{cdata(article.meta_description)}</wp:meta_value>",
            "      </wp:postmeta>",
            "    </item>",
        ))
    lines.extend(("  </channel>", "</rss>", ""))
    output.write_text("\n".join(lines), encoding="utf-8")


def write_review(articles: list[Article], review_path: Path) -> None:
    with review_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=(
            "article_id", "title", "category", "reason", "focus_keyword", "seo_title",
            "seo_title_chars", "meta_description_chars", "seo_warnings", "source", "images",
        ))
        writer.writeheader()
        for article in articles:
            writer.writerow({
                "article_id": article.article_id,
                "title": article.title,
                "category": article.category,
                "reason": article.category_reason,
                "focus_keyword": article.focus_keyword,
                "seo_title": article.seo_title,
                "seo_title_chars": len(article.seo_title),
                "meta_description_chars": len(article.meta_description),
                "seo_warnings": "; ".join(article.seo_warnings),
                "source": article.source,
                "images": article.image_count,
            })


def validate_wxr(path: Path, require_intro_images: bool = False) -> list[str]:
    from xml.etree import ElementTree as etree

    namespace = {"content": "http://purl.org/rss/1.0/modules/content/", "wp": "http://wordpress.org/export/1.2/"}
    root = etree.parse(path).getroot()
    errors: list[str] = []
    items = root.findall("./channel/item")
    if not items:
        errors.append("No post items found")
    post_names: set[str] = set()
    guids: set[str] = set()
    for index, item in enumerate(items, 1):
        if not clean_text(item.findtext("title") or ""):
            errors.append(f"Item {index} has an empty title")
        if item.findtext("wp:status", namespaces=namespace) != "draft":
            errors.append(f"Item {index} is not a draft")
        post_name = clean_text(item.findtext("wp:post_name", namespaces=namespace) or "")
        if not post_name:
            errors.append(f"Item {index} has an empty slug")
        elif post_name in post_names:
            errors.append(f"Duplicate post slug: {post_name}")
        post_names.add(post_name)
        guid = clean_text(item.findtext("guid") or "")
        if not guid:
            errors.append(f"Item {index} has an empty GUID")
        elif guid in guids:
            errors.append(f"Duplicate GUID: {guid}")
        guids.add(guid)
        metadata = {clean_text(node.findtext("wp:meta_key", namespaces=namespace) or ""): clean_text(node.findtext("wp:meta_value", namespaces=namespace) or "") for node in item.findall("wp:postmeta", namespace)}
        required = {"_yoast_wpseo_focuskw", "_yoast_wpseo_title", "_yoast_wpseo_metadesc"}
        missing = {key for key in required if not metadata.get(key)}
        if missing:
            errors.append(f"Item {index} is missing nonempty Yoast fields: {', '.join(sorted(missing))}")
        content_html = item.findtext("content:encoded", namespaces=namespace) or ""
        if not content_html:
            errors.append(f"Item {index} has an empty body")
        elif require_intro_images:
            image = re.match(r'^\s*<p><img\s+src="https?://[^"<>]+"\s+alt="[^"]*"\s+loading="lazy"\s*/></p>\s*\n', content_html)
            if not image:
                errors.append(f"Item {index} is missing a valid intro image before the body")
        categories = item.findall("category")
        if not any(node.get("domain") == "category" and clean_text(node.text or "") for node in categories):
            errors.append(f"Item {index} has no category")
        if not any(node.get("domain") == "post_tag" and clean_text(node.text or "") for node in categories):
            errors.append(f"Item {index} has no tag")
    return errors

def load_overrides(path: str | None) -> dict[str, str]:
    if not path:
        return {}
    payload = json.loads(Path(path).read_text(encoding="utf-8"))
    if not isinstance(payload, dict) or not all(value in {"行业知识", "解决方案"} for value in payload.values()):
        raise ValueError("分类覆盖文件必须是文章 ID 到 行业知识/解决方案 的 JSON 对象")
    return {str(key): str(value) for key, value in payload.items()}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", help=".docx、.zip 或包含它们的文件夹")
    parser.add_argument("--output", help="输出 WXR/XML 文件")
    parser.add_argument("--validate-only", help="只校验已有 WXR/XML 文件")
    parser.add_argument("--category-overrides", help="人工确认分类的 JSON 文件")
    parser.add_argument("--competitor-list", help="JSON competitor list used by the content audit")
    parser.add_argument("--intro-image-pool", help="JSON URL array, or an object with an images array, for intro images")
    parser.add_argument("--use-style3d-frame-pool", action="store_true", help="Use the standard 53 Style3D Frame media URLs")
    parser.add_argument("--intro-image-state", help="Persistent JSON state used to prevent image repeats within a cycle")
    parser.add_argument("--intro-image-report", help="Optional CSV report for the article-to-image assignments")
    parser.add_argument("--require-intro-images", action="store_true", help="Require one valid intro image at the start of every body when validating")
    parser.add_argument("--site-url", default="https://help-zh.style3d.com", help="WordPress 站点 URL")
    parser.add_argument("--author-login", default="weichanghua")
    parser.add_argument("--author-id", default="3")
    parser.add_argument("--brand-suffix", default="Style3D 凌迪科技")
    parser.add_argument("--dry-run", action="store_true", help="只生成分类结果，不写 XML")
    parser.add_argument("--fail-on-seo-warnings", action="store_true", help="Block XML output when the SEO review has warnings")
    args = parser.parse_args()
    if args.validate_only:
        return args
    if not args.input or not args.output:
        parser.error("生成 XML 时必须提供 --input 和 --output")
    if (args.intro_image_pool or args.use_style3d_frame_pool) and not args.intro_image_state:
        parser.error("--intro-image-state is required when intro images are enabled")
    if args.intro_image_state and not (args.intro_image_pool or args.use_style3d_frame_pool):
        parser.error("--intro-image-state requires --intro-image-pool or --use-style3d-frame-pool")
    return args


def main() -> int:
    args = parse_args()
    if args.validate_only:
        errors = validate_wxr(Path(args.validate_only), args.require_intro_images)
        if errors:
            print("校验失败:")
            print("\n".join(f"- {error}" for error in errors))
            return 1
        print("WXR 校验通过。")
        return 0

    try:
        overrides = load_overrides(args.category_overrides)
        sources = list(iter_sources(Path(args.input)))
        if not sources:
            raise ValueError("没有找到 .docx 文件")
        articles = [make_article(source, overrides, args.brand_suffix) for source in sources]
        audits = [audit_source(source, load_competitors(args.competitor_list)) for source in sources]
        duplicate_ids = [key for key, count in Counter(article.article_id for article in articles).items() if count > 1]
        if duplicate_ids:
            raise ValueError("文章 ID 重复: " + ", ".join(duplicate_ids))
        output = Path(args.output)
        review_path = output.with_suffix(".review.csv")
        audit_path = output.with_suffix(".content-audit.csv")
        write_review(articles, review_path)
        write_content_audit(audits, audit_path)
        print(f"Content audit: {audit_path}")
        pending = [article for article in articles if article.category == "待确认"]
        print(f"已读取 {len(articles)} 篇文章，分类清单: {review_path}")
        if pending:
            print("以下文章需要人工确认分类:")
            for article in pending:
                print(f"- {article.article_id}: {article.title}")
            return 2
        flagged_audits = [audit for audit in audits if audit.matches]
        if flagged_audits:
            print(f"Found configured competitor names in {len(flagged_audits)} article(s); no XML was created.")
            for audit in flagged_audits:
                print(f"- {audit.article_id}: {', '.join(name for name, _ in audit.matches)}")
            return 4
        warnings = [article for article in articles if article.seo_warnings]
        if warnings:
            print(f"SEO review has warnings for {len(warnings)} article(s).")
            for article in warnings:
                print(f"- {article.article_id}: {'; '.join(article.seo_warnings)}")
            if args.fail_on_seo_warnings:
                print("No XML created because --fail-on-seo-warnings was set.")
                return 3
        if args.dry_run:
            if args.intro_image_pool or args.use_style3d_frame_pool:
                pool = load_intro_image_pool(args.intro_image_pool, args.use_style3d_frame_pool)
                assignments, _ = assign_intro_images(articles, pool, Path(args.intro_image_state))
                report_path = Path(args.intro_image_report) if args.intro_image_report else output.with_suffix(".intro-images.csv")
                report_path.parent.mkdir(parents=True, exist_ok=True)
                write_intro_image_report(assignments, report_path)
                print(f"Intro image assignment (state unchanged): {report_path}")
            print("分类已确认；因 --dry-run 未生成 XML。")
            return 0
        output.parent.mkdir(parents=True, exist_ok=True)
        assignments: list[IntroImageAssignment] = []
        next_image_state: dict[str, object] | None = None
        if args.intro_image_pool or args.use_style3d_frame_pool:
            pool = load_intro_image_pool(args.intro_image_pool, args.use_style3d_frame_pool)
            assignments, next_image_state = assign_intro_images(articles, pool, Path(args.intro_image_state))
        assignment_by_article = {assignment.article_id: assignment for assignment in assignments}
        write_wxr(articles, output, args, assignment_by_article)
        errors = validate_wxr(output, bool(assignments))
        if errors:
            print("生成后校验失败:\n" + "\n".join(f"- {error}" for error in errors))
            return 1
        if assignments:
            report_path = Path(args.intro_image_report) if args.intro_image_report else output.with_suffix(".intro-images.csv")
            report_path.parent.mkdir(parents=True, exist_ok=True)
            write_intro_image_report(assignments, report_path)
            write_intro_image_state(Path(args.intro_image_state), next_image_state or {})
            print(f"Intro image assignment: {report_path}")
        image_total = sum(article.image_count for article in articles)
        print(f"已生成并校验: {output}")
        if image_total:
            print(f"注意: 源 Word 中有 {image_total} 个内嵌图片；XML 不会上传图片文件。")
        return 0
    except Exception as error:
        print(f"失败: {error}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
